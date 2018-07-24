"""
 *****************************************************************************
   FILE :           report.py

   AUTHOR :         Jiin Jeong

   DATE :           June 5, 2018

   DESCRIPTION :    Generates general iTree report for organizations.

 *****************************************************************************
"""

import pandas as pd
import matplotlib.pyplot as plt  # Imports plots.
import math  # Math functions.

from docx import Document
from docx.shared  import Pt  # Imports font typeface and size.
from docx.enum.style import WD_STYLE_TYPE  # Accesses styles attribute.

import datetime  # Imports cur time and date.


class csvData():
    """ This is a class for analyzing csv data. """
    def __init__(self, file, start):
        self._df = pd.read_csv(file, skiprows=start)  # Starts from start.

    def read_file(self):
        """ Reads file in mutable float format. """

        # Renames a weirdly named index.
        self._df.rename(index=str, columns={"PM2.5                                                     Avoided (pounds)":
                  "PM2.5 Avoided (pounds)", "PM2.5                                                      Removed (pounds)":
                  "PM2.5 Removed (pounds)"}, inplace=True)  #inplace makes this new label in place of df.

        # Changes col data into floats.
        for i in range(2, len(self._df.columns) - 1):  # Starts from CO2 Avoided column.
            cur_col = self._df[self._df.columns[i]]  # Current column by index.

            if type(cur_col[0]) == str:  # Checks if first value is string. Only goes through string values.
                cur_col = cur_col.str.replace("$","")  # Replaces dollar values.
                cur_col = cur_col.str.replace(",","")  # Replaces , values.
                cur_col = cur_col.astype(float)  # Changes to type float.
                self._df[self._df.columns[i]] = cur_col  # Replaces DF column with new current column.

        # print("read file")  # Tests if method was properly implemented.

    def row1(self):
        """ Reads first row and returns the itree disclaimer and location. """

        location = pd.read_csv("result.csv", header=None, nrows=1)  # Reads only the first row.
        return location.iloc[0, 0] + "," + location.iloc[0, 1]  # iloc is [row, col].

    def get_index(self):
        """ Returns a list of index of the DataFrame. """

        return pd.Index(list(self._df))

    def find_col(self, index):
        """ Finds the column by name of col index. """

        return self._df.loc[:, index]

    def sum_col(self, index):
        """ Finds the sum of the cols and rounds to the third decimal place. """

        return round(self.find_col(index).sum(), 3)

    def make_pp(self, data1, data2, data3, data4, n1, n2, n3, n4, name):
        """ Makes pie plot with given data."""

        series = pd.Series((data1, data2, data3, data4), index=[n1, n2, n3, n4], name=name)

        # Makes pie plot.
        pie = series.plot.pie(labels=[n1, n2, n3, n4],
                              colors=['r', 'y', 'g', 'c'],
                              autopct="$%.2f",  # Adds values to 2 decimal points.
                              fontsize=15,
                              figsize=(6,6))  # Sets figure size.
        fig = pie.get_figure()  # Gets figure.
        fig.savefig("pie.jpg")  # Saves figure as jpg file.


""" I wanted to make this a subclass of Document, but it was giving me a weird error:
    https://stackoverflow.com/questions/39874269/subclassing-a-function-in-python
class Report():
    def __init__(self, doc):
          Document.__init__(self, doc) """

class Report():
    def __init__(self, file):
        """ Initializes the report by creating a new file
            that is a copy of the given file. """

        self._doc = Document(file)

    def add_in(self, index, text, font="Normal"):
        """ Adds paragraph in given paragraph. """

        cur = self._doc.paragraphs[index]
        run = cur.add_run(text)
        cur_font = run.font

        if font == "Bold":  # Bold font.
            cur_font.bold = True

    def add_after(self, index, text, style="Normal"):
        """ Adds paragraph after given index. """

        self._doc.add_paragraph()
        cur = len(self._doc.paragraphs) - 1  # Current paragraph index.

        while cur > index:
            self._doc.paragraphs[cur] = self._doc.paragraphs[cur - 1]
            cur = cur - 1

        new_par = self._doc.paragraphs[cur + 1].insert_paragraph_before(text)
        new_par.style = self._doc.styles[style]

    def save(self, name):
        """ Saves the file. """

        self._doc.save(name)


"""/**** Reads CSV Data. ***/ """
data = csvData("treeresult.csv", 3)  # Reads file, skipping three lines.
data.read_file()  # Goes through data and converts to usable format.

# Sets Variables.
co2a = data.sum_col("CO2 Avoided (pounds)")
co2a_dol = data.sum_col("CO2 Avoided ($)")
co2s = data.sum_col("CO2 Sequestered (pounds)")
co2s_dol = data.sum_col("CO2 Sequestered ($)")
elec = data.sum_col("Electricity Saved (kWh)")
elec_dol = data.sum_col("Electricity Saved ($)")
fuel = data.sum_col("Fuel Saved (MMBtu)")
fuel_dol = data.sum_col("Fuel Saved ($)")
biomass = data.sum_col("Tree Biomass (short ton)")
rain = data.sum_col("Rainfall Interception (gallons)")
storm = data.sum_col("Stormwater Managed (gallons)")
storm_dol = data.sum_col("Stormwater Managed ($)")
o3 = data.sum_col("O3 Removed (pounds)")
no2a = data.sum_col("NO2 Avoided (pounds)")
no2r= data.sum_col("NO2 Removed (pounds)")
so2a = data.sum_col("SO2 Avoided (pounds)")
so2r = data.sum_col("SO2 Removed (pounds)")
voc = data.sum_col("VOC Avoided (pounds)")
pma = data.sum_col("PM2.5 Avoided (pounds)")
pmr = data.sum_col("PM2.5 Removed (pounds)")


"""/**** From here, creates REPORT ***/ """
# GENERAL INFO SOURCE(S) : Adapted from i-Tree Design Report by Jiin Jeong.
doc = Report("report_ex.docx")
# document.add_heading("Tree Benefit Report", 0)

# Adds current date.
now = datetime.datetime.now()
doc.add_in(1, "%d/%02d/%02d." % (now.year, now.month, now.day), "Bold")  # 2-digit str format.


"""/**** CO2 ***/ """
# These are specific results. General info is copied from the example file (no need to be recreated).
doc.add_after(4, "With your trees, you will avoid %s pounds of CO2. " % co2a +
          "This is equivalent to avoiding $%s of CO2." % co2a_dol, "List Bullet")
doc.add_after(5, "Your trees will sequester %s pounds of CO2. " % co2s +
          "This is equivalent to sequestering $%s of CO2." % co2s_dol, "List Bullet")
doc.add_after(6, "In total, you will save $%s by reducing %s of "
          % (co2a_dol + co2s_dol, co2a + co2s) + "atmospheric carbon dioxide " +
          "through CO2 sequestration and decreased energy production needs and emissions.",
          "List Bullet")


"""/**** ENERGY ***/ """
doc.add_after(12, "With your trees, you will save %s kWh of electricity. " % elec +
          "Your electricity energy savings are $%s." % elec_dol, "List Bullet")

doc.add_after(13, "With your trees, you will save %s MMBtu of fuel. " % fuel +
          "Your fuel savings are $%s." % fuel_dol, "List Bullet")


"""/**** ECO ***/ """
doc.add_after(20, "Your trees produce %s short tons of biomass." % biomass, "List Bullet")
doc.add_after(21, "Your trees intercept %s gallons of rainwater." % rain, "List Bullet")
doc.add_after(22, "Your trees manage $%s worth of %s gallons of stormwater."
              % (storm_dol, storm), "List Bullet")

"""/**** AIR POLLUTION ***/ """
doc.add_after(30, "Your trees remove %s pounds of O3, %s pounds of NO2, " % (o3, no2r) +
              "%s pounds of SO2, and %s pounds of PM2.5." % (so2r, pmr), "List Bullet")
doc.add_after(31, "Your trees avoid %s pounds of NO2, %s pounds of SO2, " % (no2a, so2a) +
              "%s pounds of VOC, and %s pounds of PM2.5." % (voc, pma), "List Bullet")

"""/**** SUMMARY ***/ """
# Adds pieplot of the breakdown of tree benefits.
# data.make_pp(co2a_dol + co2s_dol, elec_dol, fuel_dol, data.sum_col("Stormwater Managed ($)"),
#              "CO2", "Electricity", "Fuel", "Water", "Total Benefits in $")
# document.add_picture("pie.jpg")

doc.save("report.docx")
